
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create Document
doc = Document()

# Header Information
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = header.add_run("Shubham Singh\n")
run.bold = True
run.font.size = Pt(14)

header.add_run("6268326377\n")
header.add_run("subh7409@gmail.com\n")
header.add_run("https://www.linkedin.com/in/shubham-singh-ab9951a7/\n\n")

# Date
import datetime
today = datetime.date.today().strftime("%d %B %Y")
doc.add_paragraph(today)

# Recipient
recipient = doc.add_paragraph()
recipient.add_run("\nHiring Manager\n")
recipient.add_run("Infoverity\n")
# recipient.add_run("[Company Address]\n") # Removed per user preference for generic hiring manager address if specific address is unknown

# Subject
p_subject = doc.add_paragraph()
run_subject = p_subject.add_run("\nRE: Application for QA Automation & Manual Tester Role")
run_subject.bold = True

# Salutation
doc.add_paragraph("\nDear Hiring Manager,")

# Body - Paragraph 1
p1 = doc.add_paragraph(
    "I am writing to express my strong interest in the QA Automation & Manual Tester position at Infoverity. "
    "With 3 years of hands-on experience in delivering high-quality software solutions, I specialize in building robust test automation frameworks using Python and Selenium, while maintaining a rigorous approach to manual testing for complex workflows."
)

# Body - Paragraph 2 (Current Role at Insphere Solutions)
p2 = doc.add_paragraph(
    "In my current role at Insphere Solutions working on the Central and State Education Portal (Ministry of Education), I have honed my ability to ensure the stability of critical government platforms serving thousands of users. "
    "I successfully led the transition from manual to automated testing, designing a maintainable Page Object Model (POM) framework from scratch. "
    "This initiative reduced regression testing time by 40% and significantly improved release cycles."
)

# Body - Paragraph 3 (Why Good Fit)
p3 = doc.add_paragraph("Why I am a great fit for Infoverity:")

# Bullets
bullets = [
    "Automation Expertise: I have extensive experience developing scalable automation suites using Python, Selenium WebDriver, and Pytest. My frameworks feature dynamic data generation, smart wait mechanisms, and robust handling of complex elements like Shadow DOM, Payment Gateways (SabPaisa), and dynamic dropdowns.",
    "Manual Testing Proficiency: I believe automation complements, not replaces, manual testing. I am skilled in exploratory testing, functional decomposition, and creating detailed test cases that uncover edge-case defects in user registration, secure login flows, and document management systems.",
    "Problem-Solving & Adaptability: I have engineered unique solutions for \"Human-in-the-Loop\" scenarios, seamlessly integrating manual checkpoints (like OTP and CAPTCHA) into automated runs to ensure 100% test coverage of security-critical features.",
    "Collaboration: I am accustomed to working closely with developers and stakeholders, providing detailed defect reports and HTML execution logs to accelerate bug fixing and ensure alignment with business requirements."
]

for point in bullets:
    doc.add_paragraph(point, style='List Bullet')

# Closing
p_closing = doc.add_paragraph(
    "\nI am eager to bring my technical skills and dedication to quality to the team at Infoverity. "
    "I am confident that my background in both manual rigor and automation efficiency will allow me to contribute immediately to your software quality goals."
)

p_thank_you = doc.add_paragraph(
    "Thank you for your time and consideration. I look forward to the possibility of discussing how my experience aligns with your team's needs."
)

# Sign-off
sign_off = doc.add_paragraph("\nSincerely,\n\nShubham Singh")

# Save
file_path = "Shubham_Singh_Cover_Letter_Infoverity.docx"
doc.save(file_path)
print(f"Document saved to {file_path}")
